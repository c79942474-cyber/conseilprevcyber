"""Export des livrables — Markdown → Word (.docx) et PDF mis en page CONSEILPREV.

Le générateur (livrables.py + assistant.generate) produit du Markdown ; ce module
le transforme en document professionnel. Les deux formats partagent délibérément
la MÊME charte — mêmes couleurs, même bloc de garde, mêmes tableaux, même pied de
page — pour qu'un livrable soit reconnaissable quel que soit le format choisi :
un client qui reçoit le PDF et le Word ne doit pas croire à deux documents.

python-docx et fpdf2 sont déjà des dépendances ; aucun moteur externe requis.
"""
import io
import os
import re

# Le nettoyage des titres de documents : un dépôt sert des noms de
# fichiers, un livrable cite des références. Une seule règle, partagée
# avec les composeurs de documents.
import extraits as _X

HERE = os.path.dirname(os.path.abspath(__file__))
EMBLEM = os.path.join(HERE, "emblem.png")

# --- Charte commune aux deux formats ------------------------------------------
# Une seule définition, déclinée en RGBColor (Word) ou en tuple (PDF) au moment
# de l'usage : la couleur d'un titre ne peut pas diverger d'un format à l'autre.
C_NAVY = (0x0A, 0x22, 0x30)      # titres, texte fort
C_TEAL = (0x0E, 0x6D, 0x7C)      # accent de marque
C_GREY = (0x55, 0x66, 0x66)      # mentions secondaires
C_LINE = (0xCB, 0xD5, 0xDB)      # filets et bordures
C_BAND = (0xEC, 0xF3, 0xF5)      # fond des en-têtes de tableau et du bloc de garde
C_ZEBRA = (0xF7, 0xFA, 0xFB)     # fond des lignes paires

CONTACT = ("CONSEILPREV · christophe.cerf@outlook.com · +33 6 60 69 21 45 · "
           "conseilprevcyber.onrender.com")
MENTION = ("Brouillon généré avec l'aide de l'IA à partir de la base de connaissance "
           "CONSEILPREV — à relire, compléter et valider par un consultant.")


def _hex(rgb):
    return "%02X%02X%02X" % rgb


# --- Typographie française ---------------------------------------------------
# Appliquée AU RENDU et non dans les sources : les modules écrivent des chaînes
# Python, où l'apostrophe droite est la plus sûre, et une règle posée ici couvre
# tout ce qui sort — y compris le texte produit par le modèle, qu'on ne relit
# pas ligne à ligne. Corriger les 493 apostrophes du référentiel à la main
# aurait laissé passer la 494e, écrite demain.
_LIEN_MD = re.compile(r"\[[^\]\n]*\]\([^)\s]*\)|https?://\S+")
_APOS = re.compile(r"(?<=\w)'(?=\w)")
# L'espace fine insécable devant la ponctuation haute. U+202F : présente dans
# Liberation Sans, et c'est la raison pour laquelle cette police est embarquée
# plutôt qu'une autre.
_FINE = " "
_PONCT = re.compile(r"[  ]?([;:!?%])(?=\s|$)")


def typographie(s):
    """Applique les usages français au texte d'un livrable.

    Les adresses et les liens Markdown sont mis à l'abri d'abord : « http:// »
    ne prend pas d'espace avant ses deux-points, et une apostrophe dans une URL
    n'est pas une apostrophe de texte.
    """
    if not s:
        return s
    abris, out, pos = [], [], 0
    for m in _LIEN_MD.finditer(s):
        out.append(s[pos:m.start()])
        out.append("\x00%d\x00" % len(abris))
        abris.append(m.group(0))
        pos = m.end()
    out.append(s[pos:])
    t = "".join(out)
    t = _APOS.sub("’", t)
    # Le pourcentage ne prend la fine que s'il suit un nombre : « 100 % » oui,
    # « %s » non — ce dernier n'est pas un pourcentage mais un format.
    t = re.sub(r"(?<=\d)[  ]?%", _FINE + "%", t)
    t = _PONCT.sub(lambda m: (_FINE + m.group(1)) if m.group(1) != "%"
                   else m.group(0), t)
    for i, v in enumerate(abris):
        t = t.replace("\x00%d\x00" % i, v)
    return t


# --- Analyse Markdown en blocs -----------------------------------------------
_INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*\n]+\*)")


def _blocks(md):
    # La typographie est appliquée ICI, en amont du découpage : c'est le seul
    # point par lequel passent les deux formats. La poser dans chacun d'eux
    # aurait produit un Word et un PDF composés différemment — exactement ce
    # que ce module existe pour éviter.
    lines = typographie(md or "").replace("\r", "").split("\n")
    out, i, n = [], 0, len(lines)
    while i < n:
        ln = lines[i]
        if not ln.strip():
            i += 1
            continue
        h = re.match(r"^(#{1,6})\s+(.*)$", ln)
        if h:
            out.append(("h" + str(min(len(h.group(1)), 3)), h.group(2).strip()))
            i += 1
            continue
        if re.match(r"^\s*([-*_])\1{2,}\s*$", ln):
            out.append(("hr", None))
            i += 1
            continue
        # tableau : ligne d'en-tête + ligne de séparation
        if (re.match(r"^\s*\|.*\|\s*$", ln) and i + 1 < n
                and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]) and "-" in lines[i + 1]):
            head = [c.strip() for c in ln.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < n and re.match(r"^\s*\|.*\|\s*$", lines[i]):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            out.append(("table", (head, rows)))
            continue
        # CITATION. Le marqueur « > » n'était pas reconnu : il sortait tel quel
        # dans le document remis, collé au texte cité et répété à chaque ligne
        # (« > luer l'état actuel  > de l'empreinte… »). Nos documents s'en
        # servent pour l'avertissement de tête ET pour les extraits reproduits
        # mot pour mot — c'est-à-dire aux deux endroits où l'on doit voir d'un
        # coup d'œil que le texte n'est pas de nous.
        if re.match(r"^\s*>\s?", ln):
            cite = []
            while i < n and re.match(r"^\s*>\s?", lines[i]):
                cite.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            # Les lignes d'une même citation forment un paragraphe ; une ligne
            # vide (« > » seul) en ouvre un nouveau.
            paras, cur = [], []
            for c in cite:
                if c.strip():
                    cur.append(c.strip())
                elif cur:
                    paras.append(" ".join(cur))
                    cur = []
            if cur:
                paras.append(" ".join(cur))
            if paras:
                out.append(("quote", paras))
            continue
        if re.match(r"^\s*[-*]\s+", ln):
            # LE NIVEAU DE LA PUCE est retenu. Sans lui, un sommaire à deux
            # étages ressortait à plat : « 1. Objet » et « 1.1 Nature » au même
            # rang, ce qui défait précisément ce qu'un sommaire sert à montrer.
            items = []
            while i < n and re.match(r"^\s*[-*]\s+", lines[i]):
                creux = len(lines[i]) - len(lines[i].lstrip(" "))
                items.append((min(creux // 2, 2),
                              re.sub(r"^\s*[-*]\s+", "", lines[i])))
                i += 1
            out.append(("ul", items))
            continue
        if re.match(r"^\s*\d+[.)]\s+", ln):
            items = []
            while i < n and re.match(r"^\s*\d+[.)]\s+", lines[i]):
                creux = len(lines[i]) - len(lines[i].lstrip(" "))
                items.append((min(creux // 2, 2),
                              re.sub(r"^\s*\d+[.)]\s+", "", lines[i])))
                i += 1
            out.append(("ol", items))
            continue
        para = [ln]
        i += 1
        while (i < n and lines[i].strip()
               and not re.match(r"^(#{1,6}\s|\s*[-*]\s|\s*\d+[.)]\s|\s*\||\s*>)",
                                lines[i])):
            para.append(lines[i])
            i += 1
        out.append(("p", " ".join(para)))
    return out


def _fiche(meta):
    """Lignes du bloc de garde (« document control »), champs vides écartés.

    Ces informations existaient déjà côté application mais n'atteignaient aucun
    des deux formats : `meta` était accepté puis ignoré. Le lecteur n'avait donc
    ni destinataire, ni périmètre, ni date sous les yeux."""
    meta = meta or {}
    # L'ORDRE EST CELUI D'UN CARTOUCHE. On identifie le document (numéro,
    # objet), on dit à quel moment du projet il se situe (phase, indice, date),
    # puis pour qui (client, périmètre), et enfin ce qu'il vaut (statut). Ces
    # trois-là manquaient : sans numéro ni indice, deux versions du même
    # document se ressemblent à s'y méprendre une fois imprimées, et sans la
    # phase on ne sait pas de quel jalon elles répondent.
    champs = [("Numéro du document", meta.get("numero")),
              ("Objet", meta.get("label")),
              ("Phase du projet", meta.get("phase")),
              ("Indice", meta.get("indice")),
              ("Date d'émission", meta.get("date")),
              ("Discipline", meta.get("discipline")),
              ("Émetteur", meta.get("emetteur")),
              ("Client / organisation", meta.get("client")),
              ("Secteur d'activité", meta.get("secteur")),
              ("Périmètre", meta.get("perimetre")),
              ("Établi par", meta.get("model")),
              # Le statut RÉEL quand l'appelant le connaît. Écrire « brouillon »
              # sur un document visé lui ôterait la valeur qu'il vient
              # d'acquérir ; l'inverse serait pire.
              ("Statut", meta.get("statut") or "Brouillon, à relire et valider")]
    return [(k, str(v).strip()) for k, v in champs if v and str(v).strip()]


def _mention_pied(meta):
    """La mention du pied de page — elle suit le STATUT RÉEL du document.

    Elle était écrite en dur : « Brouillon à valider », sur toutes les pages,
    dans les deux formats. Un document VISÉ s'exportait donc en contredisant
    son propre cartouche, page après page — et c'est le pied de page qu'on lit
    en diagonale quand on feuillette un tirage, pas la page de garde. Un
    livrable validé qui circule avec « brouillon » imprimé partout se fait
    renvoyer ; l'inverse — un brouillon qui ne le dit plus — serait pire.

    Forme courte : le pied de page n'a pas la place d'une phrase. « Référence
    interne — à confronter aux textes en vigueur » y entre comme « Référence
    interne », et le cartouche porte la version complète.
    """
    s = str((meta or {}).get("statut") or "").strip()
    if not s:
        s = "Brouillon à valider"
    s = re.split(r"\s+[—–-]\s+|\s*[;(]", s)[0].strip(" .,")
    return "CONSEILPREV Cyber · " + (s[:46].rstrip() if len(s) > 46 else s)


def _sources(meta):
    """Documents de la base effectivement mobilisés, dédoublonnés.

    Le titre est celui d'une RÉFÉRENCE, pas un chemin de disque : un dépôt
    sert des noms de fichiers, et « Numeum___Contribution___Consultation_
    Arcep_… .pdf » se lit de loin comme un pointillé au bas d'un livrable
    remis. Le nettoyage est celui du module partagé — une seule règle pour
    l'annexe des sources, les attributions sous les citations et la liste
    rendue à l'écran.
    """
    out, vus = [], set()
    for s in ((meta or {}).get("sources") or []):
        titre = _X.titre_document((s.get("title") or "").strip())
        if not titre or titre in vus:
            continue
        vus.add(titre)
        out.append((titre, (s.get("theme") or "").strip(),
                    "interne" if s.get("visibility") == "internal" else "public"))
    return out


# =============================================================================
#  Word (.docx)
# =============================================================================

def _lien_docx(paragraph, texte, url, color=None):
    """Un vrai lien hypertexte Word, pas un texte bleu souligné.

    python-docx n'expose pas les liens : il faut déclarer la relation dans la
    partie du document et poser un w:hyperlink. Sans cela, les renvois entre
    livrables ne seraient cliquables qu'en PDF — le Word afficherait la syntaxe
    Markdown brute, ce qui est pire que pas de lien du tout.
    """
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    lien = OxmlElement("w:hyperlink")
    lien.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), _hex(C_TEAL))
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(c)
    rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = texte
    r.append(t)
    lien.append(r)
    paragraph._p.append(lien)
    return lien


# Le lien Markdown, reconnu AVANT le gras et l'italique : son libellé peut en
# contenir, et découper d'abord sur les astérisques casserait l'adresse.
_LIEN = re.compile(r"\[([^\]\n]+)\]\((https?://[^\s)]+)\)")


def _add_runs(paragraph, text, color=None):
    """Ajoute le texte au paragraphe en interprétant les liens, **gras**,
    *italique* et `code`."""
    from docx.shared import Pt, RGBColor
    texte = text or ""
    if _LIEN.search(texte):
        pos = 0
        for m in _LIEN.finditer(texte):
            if m.start() > pos:
                _add_runs(paragraph, texte[pos:m.start()], color=color)
            _lien_docx(paragraph, m.group(1), m.group(2))
            pos = m.end()
        if pos < len(texte):
            _add_runs(paragraph, texte[pos:], color=color)
        return
    for part in _INLINE.split(texte):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            # LE MOT QUI DÉCIDE SE VOIT. Le gras seul, dans une page de gris,
            # se remarque à peine : on l'appuie de la couleur de marque et d'un
            # fond clair, comme un surlignage sobre. C'est ce que le lecteur
            # cherche quand il parcourt sans lire.
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = RGBColor(*C_NAVY)
            shd = _el("w:shd", val="clear", fill=_hex(C_BAND))
            run._r.get_or_add_rPr().append(shd)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
        if color is not None:
            run.font.color.rgb = color


def _el(tag, **attrs):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn("w:" + k), v)
    return e


def _rule(doc):
    """Filet horizontal (bordure basse d'un paragraphe vide)."""
    p = doc.add_paragraph()
    pbdr = _el("w:pBdr")
    pbdr.append(_el("w:bottom", val="single", sz="6", space="1", color=_hex(C_LINE)))
    p._p.get_or_add_pPr().append(pbdr)
    return p


def _shade(cell, rgb):
    """Fond de cellule — indispensable pour distinguer l'en-tête d'un tableau."""
    cell._tc.get_or_add_tcPr().append(
        _el("w:shd", val="clear", color="auto", fill=_hex(rgb)))


def _champ(paragraph, instr):
    """Insère un champ Word (PAGE, NUMPAGES…) : une numérotation écrite en dur
    serait fausse dès la première modification du document par le consultant."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run = paragraph.add_run()
    deb = OxmlElement("w:fldChar")
    deb.set(qn("w:fldCharType"), "begin")
    txt = OxmlElement("w:instrText")
    txt.set(qn("xml:space"), "preserve")
    txt.text = " %s " % instr
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    for e in (deb, txt, fin):
        run._r.append(e)
    return run


def _page_field(paragraph):
    """« page X / Y » via des champs Word."""
    _champ(paragraph, "PAGE")
    paragraph.add_run(" / ")
    _champ(paragraph, "NUMPAGES")


# ── Marquage IA lisible par une machine (AI Act, art. 50.2) ────────────────
# La mention imprimée en bas de page renseigne un lecteur humain ; elle ne
# survit ni à un copier-coller, ni à un traitement automatisé. Le règlement
# demande un marquage EXPLOITABLE PAR UNE MACHINE : on le place donc dans les
# propriétés du fichier, qui voyagent avec lui et se lisent sans l'ouvrir.
MARQUE_IA = "AI-generated"
MARQUE_REF = "Reglement (UE) 2024/1689, art. 50"


def _marque_ia(meta):
    """Valeurs de marquage, dérivées des métadonnées du document."""
    meta = meta or {}
    modele = str(meta.get("model") or "").strip()
    return {
        "marque": MARQUE_IA,
        "producteur": "CONSEILPREV — assistance par intelligence artificielle"
                      + (" (" + modele + ")" if modele else ""),
        "mots_cles": "%s; AI-assisted; %s; brouillon-a-valider"
                     % (MARQUE_IA, MARQUE_REF),
        "note": "Contenu produit avec l'assistance d'un systeme d'intelligence "
                "artificielle et signale comme tel au titre du %s. "
                "Brouillon soumis a relecture et validation humaine." % MARQUE_REF,
    }


def build_docx(md, meta=None):
    """Construit le document Word (bytes) à partir du Markdown du livrable."""
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    meta = meta or {}
    NAVY, TEAL, GREY = RGBColor(*C_NAVY), RGBColor(*C_TEAL), RGBColor(*C_GREY)

    doc = Document()
    # Marquage IA dans les propriétés du fichier : il survit à la copie et se
    # lit sans ouvrir le document (Explorateur, indexation, outillage).
    try:
        m = _marque_ia(meta)
        cp = doc.core_properties
        cp.title = str(meta.get("label") or "Document")
        cp.subject = m["marque"]
        cp.category = m["marque"]
        cp.keywords = m["mots_cles"]
        cp.comments = m["note"]
        cp.author = m["producteur"]
        cp.last_modified_by = m["producteur"]
    except Exception:
        # Un marquage qui échoue ne doit pas empêcher la production du document :
        # la mention visible, elle, est toujours présente.
        pass
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # --- En-tête (lettre à en-tête) ---
    if os.path.exists(EMBLEM):
        try:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.add_run().add_picture(EMBLEM, width=Inches(0.55))
        except Exception:
            pass
    brand = doc.add_paragraph()
    brand.paragraph_format.space_after = Pt(0)
    r = brand.add_run("CONSEILPREV ")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = NAVY
    r2 = brand.add_run("Cyber")
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = TEAL
    sp = doc.add_paragraph()
    sub = sp.add_run("Cybersécurité industrielle IT / OT / IIoT")
    sub.italic = True
    sub.font.size = Pt(8.5)
    sub.font.color.rgb = GREY
    _rule(doc)

    # --- Bloc de garde ---
    fiche = _fiche(meta)
    if fiche:
        t = doc.add_table(rows=0, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for k, v in fiche:
            cells = t.add_row().cells
            cells[0].width = Inches(1.7)
            cells[1].width = Inches(4.6)
            rk = cells[0].paragraphs[0].add_run(k)
            rk.bold = True
            rk.font.size = Pt(9)
            rk.font.color.rgb = NAVY
            rv = cells[1].paragraphs[0].add_run(v)
            rv.font.size = Pt(9)
            for c in cells:
                c.paragraphs[0].paragraph_format.space_after = Pt(2)
            _shade(cells[0], C_BAND)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # --- Corps ---
    for kind, payload in _blocks(md):
        if kind in ("h1", "h2", "h3"):
            level = {"h1": 0, "h2": 1, "h3": 2}[kind]
            heading = doc.add_heading(level=level)
            heading.paragraph_format.space_before = Pt(14 if level else 6)
            heading.paragraph_format.space_after = Pt(4)
            _add_runs(heading, payload, color=NAVY)
            if level == 1:                       # filet sous les titres de section
                pbdr = _el("w:pBdr")
                pbdr.append(_el("w:bottom", val="single", sz="6", space="4",
                                color=_hex(C_LINE)))
                heading._p.get_or_add_pPr().append(pbdr)
        elif kind == "p":
            p = doc.add_paragraph()
            # ALIGNÉ À GAUCHE. Justifié, le texte étirait les lignes portant un
            # code de pièce — « SPC-SSI (Spécification technique — prévention
            # et sécurité incendie) » ne se coupe pas — et ouvrait des rivières
            # blanches au milieu du paragraphe.
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(7)
            _add_runs(p, payload)
        elif kind == "quote":
            # Retrait ET filet à gauche : le lecteur doit voir sans lire que ce
            # texte n'est pas de nous. Un simple italique ne suffit pas — le
            # document en emploie déjà pour ses propres notes.
            for k, bout in enumerate(payload):
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.left_indent = Inches(0.32)
                pf.space_before = Pt(6 if k == 0 else 2)
                pf.space_after = Pt(6 if k == len(payload) - 1 else 2)
                pbdr = _el("w:pBdr")
                pbdr.append(_el("w:left", val="single", sz="18", space="10",
                                color=_hex(C_TEAL)))
                p._p.get_or_add_pPr().append(pbdr)
                _add_runs(p, bout, color=GREY)
        elif kind in ("ul", "ol"):
            base = "List Bullet" if kind == "ul" else "List Number"
            for niveau, it in payload:
                # Word nomme ses niveaux « List Bullet 2 », « List Bullet 3 ».
                # Un style absent du modèle lèverait : on retombe alors sur le
                # niveau 1 plutôt que de perdre la liste.
                style = base if not niveau else "%s %d" % (base, niveau + 1)
                try:
                    p = doc.add_paragraph(style=style)
                except KeyError:
                    p = doc.add_paragraph(style=base)
                _add_runs(p, it)
        elif kind == "table":
            head, rows = payload
            cols = max(1, len(head))
            table = doc.add_table(rows=1, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j in range(cols):
                cell = table.rows[0].cells[j]
                run = cell.paragraphs[0].add_run(head[j] if j < len(head) else "")
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _shade(cell, C_NAVY)
            for k, row in enumerate(rows):
                cells = table.add_row().cells
                for j in range(cols):
                    _add_runs(cells[j].paragraphs[0], row[j] if j < len(row) else "")
                    for r_ in cells[j].paragraphs[0].runs:
                        r_.font.size = Pt(9.5)
                    if k % 2:                     # zébrure : lecture des lignes longues
                        _shade(cells[j], C_ZEBRA)
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
        elif kind == "hr":
            _rule(doc)

    # --- Annexe : sources mobilisées ---
    srcs = _sources(meta)
    if srcs:
        h = doc.add_heading(level=1)
        h.paragraph_format.space_before = Pt(16)
        _add_runs(h, "Sources mobilisées", color=NAVY)
        intro = doc.add_paragraph()
        ri = intro.add_run("Extraits de la base de connaissance CONSEILPREV ayant "
                           "servi à la rédaction de ce brouillon.")
        ri.italic = True
        ri.font.size = Pt(9)
        ri.font.color.rgb = GREY
        for titre, theme, vis in srcs:
            p = doc.add_paragraph(style="List Bullet")
            rt = p.add_run(titre)
            rt.bold = True
            rt.font.size = Pt(9.5)
            reste = " — %s · %s" % (theme or "sans thème", vis)
            rr = p.add_run(reste)
            rr.font.size = Pt(9)
            rr.font.color.rgb = GREY

    # --- Pied de page (toutes les pages) ---
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fr = fp.add_run(_mention_pied(meta) + " · page ")
    fr.font.size = Pt(8)
    fr.font.color.rgb = GREY
    _page_field(fp)
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GREY

    _rule(doc)
    note = doc.add_paragraph().add_run(typographie(MENTION))
    note.italic = True
    note.font.size = Pt(8.5)
    note.font.color.rgb = GREY
    contact = doc.add_paragraph().add_run(typographie(CONTACT))
    contact.font.size = Pt(8.5)
    contact.font.color.rgb = GREY

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =============================================================================
#  PDF (fpdf2, sans dépendance système)
# =============================================================================
# Les polices de base sont encodées en Latin-1 : on translittère les quelques
# caractères hors jeu (tirets longs, flèches, guillemets courbes, emoji…) pour un
# rendu fiable, tout en conservant les accents français.
_PDF_MAP = {
    "—": " - ", "–": "-", "−": "-", "→": "->", "←": "<-",
    "⇒": "=>", "≤": "<=", "≥": ">=", "…": "...", "•": "·",
    "▪": "·", "‘": "'", "’": "'", "“": '"', "”": '"',
    " ": " ", " ": " ", "‹": "<", "›": ">", "✓": "v",
    "œ": "oe", "Œ": "OE", "€": "EUR",
}
_INLINE_STRIP = re.compile(r"\*\*([^*]+)\*\*|`([^`]+)`|\*([^*\n]+)\*")

# --- Police embarquée --------------------------------------------------------
# Les polices de base du PDF sont limitées au Latin-1 : elles REFUSENT la
# ligature œ. Le contournement — l'écrire « oe » — produisait « Maîtrise
# d'oeuvre » dans un document remis à un client, c'est-à-dire une faute
# d'orthographe. Une police Unicode est donc embarquée, et livrée avec le dépôt
# plutôt que cherchée sur la machine : un export qui dépend des polices
# installées ne rend pas le même document ici et sur le serveur.
_FONTS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")
POLICE = "LiberationSans"
_POLICE_FICHIERS = {
    "": os.path.join(_FONTS, "LiberationSans-Regular.ttf"),
    "B": os.path.join(_FONTS, "LiberationSans-Bold.ttf"),
    "I": os.path.join(_FONTS, "LiberationSans-Italic.ttf"),
    # QUATRIÈME GRAISSE. fpdf la réclame dès qu'un texte porte du gras DANS de
    # l'italique — « *voir **SPC-SSI**, à la même phase* ». Sans elle, l'export
    # ne dégradait pas : il levait « Undefined font: liberationsansBI » et ne
    # rendait AUCUN PDF. Un livrable entier perdu pour une combinaison de
    # style que tout rédacteur, humain ou modèle, finit par écrire.
    "BI": os.path.join(_FONTS, "LiberationSans-BoldItalic.ttf"),
}
# Les graisses sans lesquelles le document perdrait une distinction qu'il
# utilise. Le gras-italique n'en fait pas partie : il se remplace par le gras
# sans que le lecteur perde une information.
_POLICE_REQUISES = ("", "B", "I")


def police_unicode_disponible():
    """Les graisses nécessaires sont-elles présentes ? Pour le contrôle de santé.

    Il en faut TROIS : sans l'italique, fpdf retomberait silencieusement sur le
    romain et le document perdrait une distinction qu'il utilise (les mentions
    de niveau et les notes sont en italique).
    """
    return all(os.path.exists(_POLICE_FICHIERS[s]) for s in _POLICE_REQUISES)


def _enregistrer_police(pdf):
    """Enregistre la police Unicode. Renvoie la famille à utiliser.

    Si les fichiers manquent, on retombe sur Helvetica ET sur la
    translittération — le document reste produit, avec « oe » pour « œ », ce
    qui est un défaut connu et non une panne. Le contrôle de santé signale le
    cas ; il ne doit pas se découvrir sur un livrable client.

    Le gras-italique absent, lui, ne coûte rien au lecteur : on l'écrit avec le
    fichier gras. Refuser l'export pour cette seule graisse coûterait le
    document entier.
    """
    if not police_unicode_disponible():
        return "Helvetica", False
    try:
        for style, chemin in _POLICE_FICHIERS.items():
            if not os.path.exists(chemin):
                chemin = _POLICE_FICHIERS["B" if "B" in style else ""]
            pdf.add_font(POLICE, style, chemin)
        return POLICE, True
    except Exception:
        return "Helvetica", False


def _pdf_txt(s, unicode_ok=True):
    """Prépare un texte pour le PDF.

    Avec la police Unicode, on ne translittère plus rien : le texte part tel
    qu'il a été écrit. Sans elle, on retombe sur l'ancien comportement.

    Les marqueurs Markdown ne sont PAS retirés ici : fpdf2 sait les rendre
    (`markdown=True`), et les retirer revenait à produire un PDF sans gras là
    où le Word en avait — le même livrable, deux documents différents.
    """
    s = s or ""
    if unicode_ok:
        # fpdf2 note l'italique __ainsi__ ; notre Markdown l'écrit *ainsi*. Sans
        # conversion, les astérisques s'imprimeraient tels quels. Le gras
        # (**…**) et les liens ([…](…)) ont la même notation des deux côtés.
        # Le gras est mis à l'abri d'abord : sinon la règle de l'italique
        # découperait ses deux astérisques.
        garde = "\x00\x00"
        s = s.replace("**", garde)
        s = re.sub(r"\*([^*\n]+)\*", r"__\1__", s)
        s = s.replace(garde, "**")
        # Le code en ligne : fpdf ne sait pas le rendre, et laisser les accents
        # graves ferait croire à une coquille. On retire les marqueurs.
        s = re.sub(r"`([^`]+)`", r"\1", s)
        return s
    s = _INLINE_STRIP.sub(lambda m: m.group(1) or m.group(2) or m.group(3) or "", s)
    for k, v in _PDF_MAP.items():
        s = s.replace(k, v)
    # Les substitutions encadrées d'espaces (« — » -> « - ») en produisent en
    # double ; visibles à l'impression, ils font négligé.
    s = re.sub(r" {2,}", " ", s)
    return s.encode("latin-1", "ignore").decode("latin-1")


def _pdf_class():
    """Classe PDF avec en-tête courant et pied de page numéroté.

    Définie dans une fonction : fpdf n'est importé qu'au moment de l'export,
    l'application démarre donc même si la bibliothèque manque."""
    from fpdf import FPDF

    class _Livrable(FPDF):
        titre_courant = ""
        # Résolues par build_pdf. En attributs de classe pour que l'en-tête et
        # le pied de page, appelés par fpdf sans argument, y aient accès.
        police = "Helvetica"
        uni = False
        mention = "CONSEILPREV Cyber · Brouillon à valider"

        def header(self):
            # Page 1 : c'est la lettre à en-tête qui tient ce rôle.
            if self.page_no() == 1:
                return
            self.set_font(self.police, "", 8)
            self.set_text_color(*C_GREY)
            self.cell(0, 5, _pdf_txt(self.titre_courant, self.uni)[:110],
                      new_x="LMARGIN", new_y="NEXT")
            self.set_draw_color(*C_LINE)
            y = self.get_y()
            self.line(self.l_margin, y, self.w - self.r_margin, y)
            self.ln(3)
            self.set_text_color(0, 0, 0)

        def footer(self):
            self.set_y(-13)
            self.set_font(self.police, "", 8)
            self.set_text_color(*C_GREY)
            self.cell(self.epw / 2, 5, _pdf_txt(self.mention, self.uni), align="L")
            self.cell(self.epw / 2, 5, "page %d / {nb}" % self.page_no(), align="R")

    return _Livrable


def build_pdf(md, meta=None):
    """Construit le document PDF (bytes) à partir du Markdown du livrable."""
    meta = meta or {}
    pdf = _pdf_class()(format="A4", unit="mm")
    # Même marquage que pour le Word : les métadonnées PDF sont lisibles par
    # une machine sans rendu de la page.
    try:
        m = _marque_ia(meta)
        pdf.set_title(str(meta.get("label") or "Document"))
        pdf.set_author(m["producteur"])
        pdf.set_subject(m["marque"])
        pdf.set_keywords(m["mots_cles"])
        pdf.set_creator(m["producteur"])
    except Exception:
        pass
    pdf.titre_courant = (meta.get("label") or "Livrable") + (
        " — " + meta["client"] if meta.get("client") else "")
    # Le pied de page est appelé par fpdf sans argument : la mention doit y
    # être posée AVANT la première page.
    pdf.mention = _mention_pied(meta)
    pdf.set_auto_page_break(True, margin=20)
    pdf.set_margins(16, 14, 16)
    # La police AVANT la première page : l'en-tête courant s'en sert dès la
    # page 2, et fpdf l'appelle sans passer par notre code.
    FAM, UNI = _enregistrer_police(pdf)
    pdf.police, pdf.uni = FAM, UNI
    pdf.add_page()

    # --- En-tête (lettre à en-tête) ---
    if os.path.exists(EMBLEM):
        try:
            pdf.image(EMBLEM, x=16, y=12, w=9)
            pdf.set_x(28)
        except Exception:
            pass
    pdf.set_font(FAM, "B", 15)
    pdf.set_text_color(*C_NAVY)
    pdf.cell(pdf.get_string_width("CONSEILPREV "), 8, "CONSEILPREV ",
             new_x="RIGHT", new_y="TOP")
    pdf.set_font(FAM, "B", 11)
    pdf.set_text_color(*C_TEAL)
    pdf.cell(0, 8, "Cyber", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(FAM, "I", 8.5)
    pdf.set_text_color(*C_GREY)
    pdf.cell(0, 5, _pdf_txt("Cybersécurité industrielle IT / OT / IIoT", UNI),
             new_x="LMARGIN", new_y="NEXT")
    pdf.ln(1)
    pdf.set_draw_color(*C_LINE)
    y = pdf.get_y()
    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
    pdf.ln(3)

    def rule():
        pdf.ln(1)
        pdf.set_draw_color(*C_LINE)
        yy = pdf.get_y()
        pdf.line(pdf.l_margin, yy, pdf.w - pdf.r_margin, yy)
        pdf.ln(2)

    def _cell(text, h, width_off=0.0, align="L", md_on=None):
        """multi_cell robuste : x réinitialisé, largeur explicite (jamais nulle).

        `markdown=True` fait rendre à fpdf le **gras**, l'__italique__ et les
        liens [texte](url) — cliquables dans le PDF. Sans lui, le PDF perdait
        le gras que le Word rendait : le même livrable donnait deux documents.
        Réservé au mode Unicode : avec la police de base, les marqueurs sont
        retirés en amont et il n'y a plus rien à interpréter.
        """
        pdf.set_x(pdf.l_margin + width_off)
        pdf.multi_cell(pdf.epw - width_off, h, text, align=align,
                       markdown=UNI if md_on is None else md_on)

    # --- Bloc de garde ---
    fiche = _fiche(meta)
    if fiche:
        lab_w, val_w = 42.0, pdf.epw - 42.0
        for k, v in fiche:
            y0 = pdf.get_y()
            pdf.set_fill_color(*C_BAND)
            pdf.set_font(FAM, "B", 8.5)
            pdf.set_text_color(*C_NAVY)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(lab_w, 5.5, _pdf_txt(k, UNI), fill=True, border=0,
                           new_x="RIGHT", new_y="TOP", max_line_height=5.5)
            pdf.set_font(FAM, "", 8.5)
            pdf.set_text_color(0, 0, 0)
            pdf.set_xy(pdf.l_margin + lab_w + 2, y0)
            pdf.multi_cell(val_w - 2, 5.5, _pdf_txt(v, UNI), border=0)
            pdf.set_y(max(pdf.get_y(), y0 + 5.5))
        pdf.ln(2)

    # --- Corps ---
    for kind, payload in _blocks(md):
        if kind in ("h1", "h2", "h3"):
            size, lh = {"h1": (16, 8), "h2": (13, 7), "h3": (11, 6)}[kind]
            pdf.ln(2 if kind == "h1" else 1)
            pdf.set_font(FAM, "B", size)
            pdf.set_text_color(*C_NAVY)
            _cell(_pdf_txt(payload, UNI), lh)
            if kind == "h2":                      # filet sous les titres de section
                pdf.set_draw_color(*C_LINE)
                yy = pdf.get_y()
                pdf.line(pdf.l_margin, yy, pdf.w - pdf.r_margin, yy)
                pdf.ln(2)
            pdf.set_text_color(0, 0, 0)
        elif kind == "p":
            pdf.set_font(FAM, "", 10.5)
            # ALIGNÉ À GAUCHE, plus justifié. La justification étirait les
            # lignes portant un code de pièce — « SPC-SSI (Spécification
            # technique — prévention et sécurité incendie) » ne se coupe pas —
            # et ouvrait des rivières blanches au milieu du paragraphe.
            _cell(_pdf_txt(payload, UNI), 5.2, align="L")
            pdf.ln(1.6)
        elif kind == "quote":
            # Filet vertical + retrait, comme en Word : le même document doit
            # se lire pareil dans les deux formats.
            pdf.set_font(FAM, "I", 10)
            pdf.set_text_color(*C_GREY)
            for bout in payload:
                y0 = pdf.get_y()
                pdf.set_x(pdf.l_margin + 5)
                _cell(_pdf_txt(bout, UNI), 5, width_off=5)
                y1 = pdf.get_y()
                pdf.set_draw_color(*C_TEAL)
                pdf.set_line_width(0.7)
                # Un filet ne se trace que si la citation n'a pas changé de
                # page : à cheval, il descendrait jusqu'au pied de page.
                if y1 > y0:
                    pdf.line(pdf.l_margin + 2, y0, pdf.l_margin + 2, y1 - 1)
                pdf.set_line_width(0.2)
                pdf.set_draw_color(*C_LINE)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font(FAM, "", 10.5)
            pdf.ln(1.5)
        elif kind in ("ul", "ol"):
            # RETRAIT PENDANT. La puce et son texte partaient dans la même
            # cellule : la deuxième ligne d'une puce revenait sous la puce, et
            # une liste de trois lignes ne se distinguait plus d'un paragraphe.
            # La puce s'écrit donc à part, et le texte dans une colonne qui
            # garde sa marge à toutes ses lignes.
            pdf.set_font(FAM, "", 10.5)
            for idx, (niveau, it) in enumerate(payload, 1):
                decal = 3.0 + 6.0 * niveau
                gouttiere = decal + 4.0
                marque = "·" if kind == "ul" else "%d." % idx
                txt = _pdf_txt(it, UNI)
                # LA PUCE NE PART PAS SANS SON TEXTE. Écrite la première, elle
                # déclenchait seule le saut de page ; le texte, replacé au Y
                # d'avant — devenu le bas de la page suivante — sautait à son
                # tour. Restait entre les deux une page dont tout le contenu
                # était un point d'énumération. On mesure donc l'élément avant
                # de l'écrire, et on tourne la page d'abord s'il n'y tient pas.
                haut = pdf.multi_cell(pdf.epw - gouttiere, 5, txt, align="L",
                                      markdown=UNI, dry_run=True,
                                      output="HEIGHT")
                # Un élément plus haut qu'une page pleine ne tiendra nulle part :
                # le faire sauter n'ouvrirait qu'une page blanche de plus.
                if haut <= pdf.h - pdf.t_margin - 20 and pdf.will_page_break(haut):
                    pdf.add_page()
                y0 = pdf.get_y()
                pdf.set_x(pdf.l_margin + decal)
                pdf.cell(4.0, 5, _pdf_txt(marque, UNI), align="L")
                pdf.set_xy(pdf.l_margin + gouttiere, y0)
                pdf.multi_cell(pdf.epw - gouttiere, 5, txt,
                               align="L", markdown=UNI)
            pdf.ln(1)
        elif kind == "table":
            head, rows = payload
            cols = max(1, len(head))
            pdf.set_font(FAM, "", 9)
            try:
                from fpdf.fonts import FontFace
                entete = FontFace(emphasis="BOLD", color=(255, 255, 255),
                                  fill_color=C_NAVY)
                with pdf.table(first_row_as_headings=True, line_height=5,
                               headings_style=entete, cell_fill_color=C_ZEBRA,
                               cell_fill_mode="ROWS", borders_layout="ALL",
                               text_align="LEFT", padding=1.4,
                               # Les codes de pièce du registre sont des liens :
                               # sans cela, la colonne afficherait la syntaxe
                               # Markdown brute au lieu du code.
                               markdown=UNI) as table:
                    hr = table.row()
                    for j in range(cols):
                        hr.cell(_pdf_txt(head[j], UNI) if j < len(head) else "")
                    for row in rows:
                        tr = table.row()
                        for j in range(cols):
                            tr.cell(_pdf_txt(row[j], UNI) if j < len(row) else "")
            except Exception:
                # Repli : jamais d'export perdu pour un tableau récalcitrant.
                for row in [head] + rows:
                    pdf.multi_cell(0, 5, _pdf_txt(" | ".join(row), UNI))
            pdf.ln(2)
        elif kind == "hr":
            rule()

    # --- Annexe : sources mobilisées ---
    srcs = _sources(meta)
    if srcs:
        pdf.ln(3)
        pdf.set_font(FAM, "B", 13)
        pdf.set_text_color(*C_NAVY)
        _cell(_pdf_txt("Sources mobilisées", UNI), 7)
        pdf.set_draw_color(*C_LINE)
        yy = pdf.get_y()
        pdf.line(pdf.l_margin, yy, pdf.w - pdf.r_margin, yy)
        pdf.ln(2)
        pdf.set_font(FAM, "I", 9)
        pdf.set_text_color(*C_GREY)
        _cell(_pdf_txt("Extraits de la base de connaissance CONSEILPREV ayant servi "
                       "à la rédaction de ce brouillon.", UNI), 4.5)
        pdf.ln(1)
        pdf.set_text_color(0, 0, 0)
        for titre, theme, vis in srcs:
            pdf.set_font(FAM, "B", 9.5)
            pdf.set_x(pdf.l_margin + 3)
            pdf.multi_cell(pdf.epw - 3, 5, _pdf_txt("  ·  " + titre, UNI))
            pdf.set_font(FAM, "", 8.5)
            pdf.set_text_color(*C_GREY)
            pdf.set_x(pdf.l_margin + 9)
            pdf.multi_cell(pdf.epw - 9, 4.5,
                           _pdf_txt("%s · %s" % (theme or "sans thème", vis), UNI))
            pdf.set_text_color(0, 0, 0)

    rule()
    pdf.set_font(FAM, "I", 8.5)
    pdf.set_text_color(*C_GREY)
    _cell(_pdf_txt(typographie(MENTION), UNI), 4.5)
    pdf.set_font(FAM, "", 8.5)
    _cell(_pdf_txt(typographie(CONTACT).replace("·", "-"), UNI), 4.5)

    out = pdf.output()
    return bytes(out)
