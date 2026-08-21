"""Emporter et transmettre — ce qui doit voyager AVEC le document.

DEUX MANQUES, ET LE SECOND EST LE PLUS COÛTEUX.

  1. UN RÉSULTAT NE S'EXPORTAIT PAS. La note de calcul, l'étude de phase, la
     pièce et la stratégie s'emportaient en Word et en PDF ; la TRAJECTOIRE DE
     DÉCARBONATION, non. Elle s'affichait, et le seul moyen de la sortir était
     de sélectionner le texte à l'écran — donc de perdre les tableaux, la
     distinction entre grandeur recevable et grandeur à produire, et les
     réserves des textes. Un résultat qui ne sort pas du site n'est pas un
     livrable.

  2. LE DOCUMENT PARTAIT SANS SON CONTEXTE. Sur la page, tout est dit : la
     phase, la tolérance, ce qui reste à produire. Le fichier, lui, part seul.
     Six semaines plus tard, aux achats, une enveloppe d'avant-projet à ±30 %
     devient un budget ; à l'exploitation, une valeur de conception devient une
     consigne ; en comité, un ordre de grandeur devient un engagement. Personne
     n'a menti — le contexte était resté sur le site.

CE QUE CES TESTS VERROUILLENT : que les cinq documents s'exportent, que le
bordereau arrive RÉELLEMENT DANS LE FICHIER (et pas seulement dans une réponse
JSON qu'on n'ouvre jamais), que rien de nominatif ne passe, et qu'un refus ne
prive jamais le client de son document.

SUR LES APOSTROPHES. Les documents produits passent par la typographie
française : « qu'il » y devient « qu’il ». Un test qui chercherait la chaîne
telle qu'elle est écrite dans le module échouerait sur la ponctuation en
croyant échouer sur le fond. On normalise donc avant de comparer.
"""
import io
import os
import re
import sys

import pytest

ICI = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ICI)

import transmission as T  # noqa: E402
from conftest import ADMIN_EMAIL, _assurer_admin  # noqa: E402

PAGES = ["datacenter.html", "ingenierie-datacenter.html",
         "strategie-durable-datacenter.html"]

# Une nature par route d'export : si l'une manque, un document sort sans
# pouvoir dire ce qu'il n'est pas.
NATURES_ATTENDUES = {"note_calcul", "trajectoire", "etude_phase", "piece",
                     "strategie_dd"}


def plat(s):
    """Apostrophes et espaces insécables ramenés à la forme du code source."""
    return (str(s).replace("’", "'").replace(" ", " ")
            .replace(" ", " "))


def lire(nom):
    with open(os.path.join(ICI, nom), encoding="utf-8") as f:
        return f.read()


@pytest.fixture
def client():
    import app
    _assurer_admin()
    app.app.config["TESTING"] = True
    c = app.app.test_client()
    with c.session_transaction() as s:
        s["user_email"] = ADMIN_EMAIL
    return c


# En-têtes minimaux : le contrôle d'origine refuse une requête sans Origin, et
# c'est voulu — ce test-là ne cherche pas à l'éprouver.
H = {"Origin": "http://localhost"}


def texte_docx(blob):
    import docx
    d = docx.Document(io.BytesIO(blob))
    t = "\n".join(p.text for p in d.paragraphs)
    for tab in d.tables:
        for row in tab.rows:
            t += "\n" + " | ".join(c.text for c in row.cells)
    return plat(t)


# ── 1. Le vocabulaire se tient ─────────────────────────────────────────────

def test_le_module_se_charge_sans_incoherence():
    assert T.sante()["problemes"] == []


def test_chaque_document_exportable_a_sa_nature():
    """Une nature manquante = un document qui sort sans pouvoir dire ce qu'il
    n'est pas — c'est-à-dire sans la seule chose qui empêche l'erreur."""
    assert NATURES_ATTENDUES <= set(T.NATURES), NATURES_ATTENDUES - set(T.NATURES)


@pytest.mark.parametrize("cle", sorted(T.NATURES))
def test_chaque_nature_nomme_les_confusions_reelles(cle):
    """Une réserve qui ne nomme pas l'erreur qu'elle prévient ne prévient
    rien : on la lit comme une formule de style et on passe."""
    n = T.NATURES[cle]
    assert len(n["n_est_pas"]) >= 2, cle
    for x in n["n_est_pas"]:
        assert len(x) >= 60, (cle, x)
    assert n["reste"], cle


@pytest.mark.parametrize("cle", sorted(T.DESTINATAIRES))
def test_chaque_fonction_dit_ce_qu_elle_ne_doit_pas_faire(cle):
    """La mise en garde seule laisse le lecteur juge de la conséquence. C'est
    l'INTERDIT nommé qui l'arrête — « ne consultez pas de fournisseur sur ces
    valeurs » agit là où « valeurs indicatives » n'agit pas."""
    d = T.DESTINATAIRES[cle]
    assert len(d["avant"]) >= 100, cle
    assert d["pas"].strip().endswith("."), cle
    assert re.match(r"^(Ne |N'|Ne l|Ne le|N’)", d["pas"]), (cle, d["pas"])


def test_les_deux_bouts_de_la_chaine_sont_couverts():
    """Un vocabulaire qui n'offrirait que des fonctions internes laisserait le
    client sans choix honnête quand il transmet à un tiers."""
    assert {"direction", "achats", "exploitation"} <= set(T.DESTINATAIRES)
    assert {"moe", "entreprise"} <= set(T.DESTINATAIRES), "les tiers aussi"


# ── 2. Rien de nominatif ───────────────────────────────────────────────────

@pytest.mark.parametrize("v", ["M. Dupont", "Mme Martin", "jean@exemple.fr",
                               "Dr Leblanc", "monsieur le directeur"])
def test_ce_qui_ressemble_a_une_personne_est_refuse(v):
    assert T.nominatif(v), v


@pytest.mark.parametrize("cle", sorted(T.DESTINATAIRES))
def test_aucune_fonction_du_vocabulaire_n_est_prise_pour_une_personne(cle):
    """Le garde ne doit pas refuser le vocabulaire qu'il est censé protéger."""
    assert not T.nominatif(cle)
    assert not T.nominatif(T.DESTINATAIRES[cle]["nom"])


def test_le_bordereau_dit_ce_qu_il_ne_porte_pas():
    b = T.bordereau("note_calcul", "achats")
    assert len(b["exclus"]) >= 3
    assert any("personne" in x.lower() for x in b["exclus"])
    assert any("copie" in x.lower() for x in b["exclus"]), (
        "l'absence de conservation doit être écrite : c'est ce que le client "
        "ne peut pas vérifier lui-même")


# ── 3. Le bordereau dit ce qu'il doit dire ─────────────────────────────────

def test_le_bordereau_porte_les_cinq_choses():
    b = T.bordereau("note_calcul", "achats",
                    {"phase": "APS", "indice": "01", "client": "ACME"})
    m = plat(b["markdown"])
    assert "Bordereau de transmission" in m
    assert "Achats et approvisionnement" in m
    assert "Ce que ce document est" in m
    assert "Ce qu'il n'est pas" in m
    assert "reste à produire" in m
    assert "revenez vers l'émetteur" in m


def test_la_reserve_arrive_avant_le_premier_chiffre():
    """Le bordereau se pose EN TÊTE. Une réserve qu'on ne rencontre qu'après
    avoir lu les chiffres arrive après la conclusion déjà tirée."""
    md, b = T.poser("# Note\n\nPUE 1,35\n", "note_calcul", "achats")
    m = plat(md)
    assert m.index("Ce qu'il n'est pas") < m.index("PUE 1,35")


def test_sans_destinataire_le_document_ne_change_pas():
    corps = "# Note\n\nPUE 1,35\n"
    md, b = T.poser(corps, "note_calcul", "")
    assert md == corps and b is None


def test_une_fonction_inconnue_ne_fait_pas_echouer_le_bordereau():
    """Un export qui échoue parce qu'une clé a bougé prive le client de son
    document — et il recommencera SANS bordereau."""
    b = T.bordereau("note_calcul", "service_des_licornes")
    assert b["markdown"], "le bordereau doit sortir quand même"
    assert any(x["champ"] == "destinataire" for x in b["refuses"])
    assert "Réserve sur ce bordereau" in plat(b["markdown"])


def test_une_nature_inconnue_se_dit_au_lieu_d_inventer():
    b = T.bordereau("document_mystere", "achats")
    assert b["markdown"] == "", (
        "sans nature, le bordereau ne peut pas dire ce que le document n'est "
        "pas — mieux vaut pas de bordereau qu'un bordereau vide de sens")
    assert any(x["champ"] == "nature" for x in b["refuses"])


# ── 4. Le bordereau arrive DANS le fichier ─────────────────────────────────

def test_le_bordereau_est_reellement_dans_le_document_produit(client):
    """LE CONTRÔLE QUI COMPTE. Tout le reste peut être vert et le bordereau
    rester dans une réponse JSON que personne n'ouvre. On ouvre le Word."""
    r = client.post("/api/datacenter/export", headers=H,
                    json={"puissance_it_kw": 5000, "pays": "FR",
                          "format": "docx", "client": "ACME",
                          "destinataire": "achats"})
    assert r.status_code == 200, r.data[:200]
    t = texte_docx(r.data)
    assert "Bordereau de transmission" in t
    assert "Ce qu'il n'est pas" in t
    assert "Achats et approvisionnement" in t
    assert "négocier" in t, "la mise en garde propre aux achats doit y être"


def test_le_destinataire_figure_aussi_au_cartouche(client):
    """Le bordereau est en page de garde, et une page de garde se détache d'un
    tirage agrafé. Le cartouche, lui, reste avec le corps."""
    r = client.post("/api/datacenter/export", headers=H,
                    json={"puissance_it_kw": 5000, "format": "docx",
                          "destinataire": "exploitation"})
    t = texte_docx(r.data)
    assert "Transmis à" in t
    assert "Exploitation et maintenance" in t


def test_sans_destinataire_le_document_reste_ce_qu_il_etait(client):
    r = client.post("/api/datacenter/export", headers=H,
                    json={"puissance_it_kw": 5000, "format": "docx"})
    assert r.status_code == 200
    t = texte_docx(r.data)
    assert "Bordereau de transmission" not in t
    assert "Transmis à" not in t


def test_un_nom_de_personne_est_ecarte_ET_le_document_part_quand_meme(client):
    """Deux exigences en une : rien de nominatif ne circule, et un refus ne
    prive pas le client de son document."""
    r = client.post("/api/datacenter/export", headers=H,
                    json={"puissance_it_kw": 5000, "format": "docx",
                          "destinataire": "M. Dupont"})
    assert r.status_code == 200
    t = texte_docx(r.data)
    assert "Dupont" not in t
    assert "Bordereau de transmission" not in t


# ── 5. La trajectoire s'exporte enfin ──────────────────────────────────────

def test_la_trajectoire_s_exporte_en_word(client):
    """DÉFAUT CORRIGÉ. Quatre documents sur cinq s'exportaient."""
    r = client.post("/api/datacenter/decarbonation/export", headers=H,
                    json={"puissance_it_kw": 5000, "pays": "FR",
                          "etape": "PERIM", "format": "docx"})
    assert r.status_code == 200, r.data[:300]
    t = texte_docx(r.data)
    assert "PERIM" in t
    assert "Ce qu'elle verrouille" in t, (
        "ce que l'étape verrouille doit précéder ce qu'elle produit : un "
        "lecteur qui l'apprend au dernier chapitre a déjà décidé")


def test_l_export_de_la_trajectoire_distingue_l_acquis_du_reste(client):
    r = client.post("/api/datacenter/decarbonation/export", headers=H,
                    json={"puissance_it_kw": 5000, "etape": "PERIM",
                          "format": "docx"})
    t = texte_docx(r.data)
    assert "Grandeurs recevables" in t or "à produire" in t


def test_la_trajectoire_accepte_aussi_un_bordereau(client):
    r = client.post("/api/datacenter/decarbonation/export", headers=H,
                    json={"puissance_it_kw": 5000, "etape": "PERIM",
                          "format": "docx", "destinataire": "hse"})
    t = texte_docx(r.data)
    assert "Bordereau de transmission" in t
    assert "réglementaire" in t, "la mise en garde propre au HSE"


def test_l_export_de_la_trajectoire_reste_ferme():
    """Le document porte le profil du projet : il n'est pas public."""
    import app
    c = app.app.test_client()
    r = c.post("/api/datacenter/decarbonation/export", headers=H,
               json={"puissance_it_kw": 5000, "etape": "PERIM"})
    assert r.status_code == 401, r.status_code


def test_une_etape_inconnue_le_dit(client):
    r = client.post("/api/datacenter/decarbonation/export", headers=H,
                    json={"puissance_it_kw": 5000, "etape": "INVENTEE"})
    assert r.status_code == 404


def test_sans_puissance_l_export_refuse_au_lieu_de_produire_du_vide(client):
    r = client.post("/api/datacenter/decarbonation/export", headers=H,
                    json={"etape": "PERIM"})
    assert r.status_code == 400


# ── 6. Le PDF aussi ────────────────────────────────────────────────────────

@pytest.mark.parametrize("route,corps", [
    ("/api/datacenter/export", {"puissance_it_kw": 5000}),
    ("/api/datacenter/decarbonation/export",
     {"puissance_it_kw": 5000, "etape": "PERIM"}),
])
def test_les_deux_formats_sortent(client, route, corps):
    """« PDF/Word » : les deux, pas l'un ou l'autre."""
    for fmt, mime in (("docx", "wordprocessingml"), ("pdf", "pdf")):
        r = client.post(route, headers=H, json=dict(corps, format=fmt))
        assert r.status_code == 200, (route, fmt, r.data[:160])
        assert mime in (r.headers.get("Content-Type") or ""), (route, fmt)
        assert len(r.data) > 3000, (route, fmt, len(r.data))


# ── 7. Le vocabulaire est servi, et les pages le portent ───────────────────

def test_la_route_sert_le_vocabulaire(client):
    r = client.get("/api/transmission")
    assert r.status_code == 200
    j = r.get_json()
    assert j["ok"]
    assert len(j["destinataires"]) == len(T.DESTINATAIRES)
    assert all(set(d) >= {"cle", "nom", "avant", "pas"} for d in j["destinataires"])


@pytest.mark.parametrize("page", PAGES)
def test_chaque_page_porte_le_module_et_sa_zone(page):
    h = lire(page)
    assert "/transmettre.js" in h, page
    assert 'class="tr-zone"' in h, page


def test_le_vocabulaire_n_est_pas_recopie_dans_les_pages():
    """Recopié trois fois, il aurait divergé au premier ajout — et c'est la
    copie de la page que le client aurait vue."""
    for page in PAGES:
        h = lire(page)
        for d in T.DESTINATAIRES.values():
            assert d["nom"] not in h, (page, d["nom"])


def test_le_selecteur_n_accepte_aucune_saisie_libre():
    """Un champ libre laisserait entrer un nom, et le document deviendrait
    nominatif sans que personne l'ait décidé."""
    js = lire("transmettre.js")
    assert "<select" in js
    assert "<input" not in js, "aucun champ de saisie dans le bloc"


def test_l_export_survit_a_l_absence_du_module():
    """Le module est à part. S'il n'a pas chargé, l'export doit partir sans
    bordereau plutôt que d'échouer : un document sans bordereau reste un
    document, un export cassé ne l'est pas."""
    for f in ("datacenter.js", "strategie-dd.js", "ingenierie-dc.js"):
        js = lire(f)
        assert "window.TRANSMETTRE && window.TRANSMETTRE.corps" in js, f
